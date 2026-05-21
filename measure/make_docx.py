# -*- coding: utf-8 -*-
"""기체 개념 진단 검사지(작성자용) → 디자인된 DOCX 생성.
그림(그래프·입자·혼합 분포)은 Pillow로 직접 렌더링해 임베드."""
import io
from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

KR = 'Malgun Gothic'
SS = 4  # supersample

# ---------------- figures ----------------
def _png(im):
    b = io.BytesIO(); im.save(b, 'PNG'); b.seek(0); return b

def _img(w, h):
    return Image.new('RGB', (w*SS, h*SS), 'white')

def graph(kind):
    w, h = 132, 108
    im = _img(w, h); d = ImageDraw.Draw(im)
    def s(v): return int(v*SS)
    ax = (55, 65, 81); col = (31, 41, 55); lw = int(2.4*SS)
    ox, oy = 22, 88
    d.line([s(ox), s(oy), s(122), s(oy)], fill=ax, width=SS)
    d.line([s(ox), s(oy), s(ox), s(14)], fill=ax, width=SS)
    if kind == 'inv':
        pts = []
        for i in range(31):
            t = i/30.0
            x = ox+8 + t*(114-ox)
            y = 20 + (1.0/(0.3+t*2.4))*46
            y = min(y, oy-3)
            pts.append((s(x), s(y)))
        d.line(pts, fill=col, width=lw, joint='curve')
    elif kind == 'prop':
        d.line([s(ox), s(oy), s(114), s(18)], fill=col, width=lw)
    elif kind == 'flat':
        d.line([s(ox+2), s(50), s(118), s(50)], fill=col, width=lw)
    elif kind == 'lin':
        d.line([s(ox+2), s(18), s(114), s(84)], fill=col, width=lw)
    return _png(im.resize((w, h), Image.LANCZOS))

def pbox(w, h, circles, gray=False):
    im = _img(w, h); d = ImageDraw.Draw(im)
    def s(v): return int(v*SS)
    if gray:
        d.rectangle([s(3), s(3), s(w-3), s(h-3)], fill=(203, 213, 225))
    d.rectangle([s(3), s(3), s(w-3), s(h-3)], outline=(55, 65, 81), width=int(1.4*SS))
    for (x, y, r, kind) in circles:
        box = [s(x-r), s(y-r), s(x+r), s(y+r)]
        if kind == 'fill':
            d.ellipse(box, fill=(55, 65, 81))
        elif kind == 'gray':
            d.ellipse(box, fill=(107, 114, 128))
        else:
            d.ellipse(box, outline=(37, 99, 235), width=int(1.8*SS))
    return _png(im.resize((w, h), Image.LANCZOS))

def mixbox(w, h, circles, divider=False):
    im = _img(w, h); d = ImageDraw.Draw(im)
    def s(v): return int(v*SS)
    d.rectangle([s(3), s(3), s(w-3), s(h-3)], outline=(55, 65, 81), width=int(1.4*SS))
    if divider:
        x = w//2
        for yy in range(8, h-6, 8):
            d.line([s(x), s(yy), s(x), s(yy+4)], fill=(156, 163, 175), width=SS)
    for (x, y, kind) in circles:
        r = 5.5
        box = [s(x-r), s(y-r), s(x+r), s(y+r)]
        if kind == 'fill':
            d.ellipse(box, fill=(107, 114, 128))
        else:
            d.ellipse(box, outline=(37, 99, 235), width=int(1.8*SS))
    return _png(im.resize((w, h), Image.LANCZOS))

def make_fig(fid):
    if fid in ('g_inv', 'g_prop', 'g_flat', 'g_lin'):
        return graph(fid[2:])
    if fid == 'p_before':
        c = [(17,18,7),(44,24,7),(30,44,7),(49,58,7),(17,74,7),(41,86,7),(24,108,7),(48,114,7)]
        return pbox(66, 132, [(x,y,r,'fill') for x,y,r in c])
    if fid == 'p_same':
        c = [(14,15,6.5),(33,17,6.5),(21,32,6.5),(36,37,6.5),(13,50,6.5),(31,54,6.5),(17,70,6.5),(35,72,6.5)]
        return pbox(50, 86, [(x,y,r,'fill') for x,y,r in c])
    if fid == 'p_small':
        c = [(14,15,3),(33,17,3),(21,32,3),(36,37,3),(13,50,3),(31,54,3),(17,70,3),(35,72,3)]
        return pbox(50, 86, [(x,y,r,'fill') for x,y,r in c])
    if fid == 'p_four':
        c = [(17,22,6.5),(34,38,6.5),(16,56,6.5),(34,70,6.5)]
        return pbox(50, 86, [(x,y,r,'fill') for x,y,r in c])
    if fid == 'p_filled':
        c = [(14,15,6.5),(33,17,6.5),(21,32,6.5),(36,37,6.5),(13,50,6.5),(31,54,6.5),(17,70,6.5),(35,72,6.5)]
        return pbox(50, 86, [(x,y,r,'fill') for x,y,r in c], gray=True)
    if fid == 'm_homog':
        c = [(15,17,'fill'),(40,15,'open'),(64,19,'fill'),(26,34,'open'),(51,36,'fill'),(70,43,'open'),
             (17,53,'fill'),(43,57,'open'),(68,62,'fill'),(22,75,'open'),(47,77,'fill'),(68,80,'open')]
        return mixbox(82, 92, c)
    if fid == 'm_layer':
        c = [(15,15,'open'),(34,13,'open'),(53,15,'open'),(70,17,'open'),(24,30,'open'),(47,30,'open'),
             (15,62,'fill'),(34,64,'fill'),(53,62,'fill'),(70,64,'fill'),(24,78,'fill'),(47,78,'fill')]
        return mixbox(82, 92, c)
    if fid == 'm_split':
        c = [(15,17,'open'),(32,17,'open'),(17,39,'open'),(32,43,'open'),(15,65,'open'),(32,68,'open'),
             (51,17,'fill'),(68,17,'fill'),(51,41,'fill'),(68,43,'fill'),(51,65,'fill'),(68,68,'fill')]
        return mixbox(82, 92, c, divider=True)
    if fid == 'm_cluster':
        c = [(19,19,'open'),(32,17,'open'),(24,30,'open'),(37,30,'open'),(19,41,'open'),(32,41,'open'),
             (51,56,'fill'),(64,54,'fill'),(56,67,'fill'),(69,67,'fill'),(53,79,'fill'),(66,79,'fill')]
        return mixbox(82, 92, c)
    raise ValueError(fid)

# ---------------- docx helpers ----------------
def set_kr(run, size=None, bold=None, color=None):
    run.font.name = KR
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:eastAsia'), KR)
    rf.set(qn('w:ascii'), KR); rf.set(qn('w:hAnsi'), KR)
    if size: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold
    if color: run.font.color.rgb = color

def shade(cell, hexcolor):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hexcolor)
    tcpr.append(shd)

def cell_text(cell, text, size=9.5, bold=False, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    r = p.add_run(text)
    set_kr(r, size=size, bold=bold, color=color)
    return p

GREEN = RGBColor(0x05, 0x96, 0x69)
RED = RGBColor(0xC0, 0x2C, 0x2C)
NAVY = RGBColor(0x1E, 0x3A, 0x8A)
MUTED = RGBColor(0x6b, 0x72, 0x80)
GREY = RGBColor(0x9c, 0xa3, 0xaf)
GREENFILL_STRONG = 'C6EBD8'
GREENFILL_SOFT = 'EAF7F0'

def fixed_table(t):
    t.autofit = False; t.allow_autofit = False
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    t._tbl.tblPr.append(layout)

def set_col_widths(t, widths_in):
    for row in t.rows:
        for i, w in enumerate(widths_in):
            row.cells[i].width = Inches(w)
    for i, w in enumerate(widths_in):
        t.columns[i].width = Inches(w)

def table_borders(t, color='D1D5DB', sz=4):
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        e = OxmlElement('w:' + edge)
        e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), str(sz)); e.set(qn('w:space'), '0'); e.set(qn('w:color'), color)
        borders.append(e)
    t._tbl.tblPr.append(borders)

def cell_margins(t, top=40, bottom=40, left=95, right=95):
    mar = OxmlElement('w:tblCellMar')
    for tag, val in (('top', top), ('bottom', bottom), ('start', left), ('end', right)):
        e = OxmlElement('w:' + tag); e.set(qn('w:w'), str(val)); e.set(qn('w:type'), 'dxa'); mar.append(e)
    t._tbl.tblPr.append(mar)

def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement('w:tblHeader'); th.set(qn('w:val'), 'true'); trPr.append(th)

def style_table(t, widths):
    fixed_table(t); table_borders(t); cell_margins(t); set_col_widths(t, widths)

def para_shade(p, hexcolor):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), hexcolor); pPr.append(shd)

def add_page_number(paragraph):
    run = paragraph.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve'); it.text = 'PAGE'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
    run._r.append(f1); run._r.append(it); run._r.append(f2)
    set_kr(run, size=9, color=GREY)

# ---------------- content ----------------
TF = [
    ('보일', '기체의 부피를 늘리면 압력도 함께 커진다.', '거짓'),
    ('보일', '기체를 압축하면 분자들이 서로 합쳐져 더 큰 분자가 된다.', '거짓'),
    ('보일', '진공은 주변의 기체를 빨아들이는 힘을 가진다.', '거짓'),
    ('돌턴', '혼합 기체의 전체 압력은 각 기체가 내는 부분 압력의 합이다.', '참'),
    ('돌턴', '혼합 기체에서는 무거운 기체가 전체 압력의 대부분을 차지한다.', '거짓'),
    ('기체 일반', '기체의 압력은 분자들이 용기 벽에 끊임없이 충돌하기 때문에 생긴다.', '참'),
    ('기체 일반', '용기의 모양이나 재질이 달라도, 온도·부피·기체의 양이 같으면 압력은 같다.', '참'),
]

def T(ang, stem, rows):  # text question
    return {'kind': 'text', 'ang': ang, 'stem': stem, 'rows': rows}

def F(ang, stem, figs, reasons, before=None):  # figure question
    return {'kind': 'fig', 'ang': ang, 'stem': stem, 'figs': figs, 'reasons': reasons, 'before': before}

SECTIONS = [
 ('오개념 B1 — 압력과 부피의 관계', [
   T('객관식 1 — 직접 판단(정량)',
     '일정한 온도에서 밀폐된 주사기의 부피를 처음의 1/3로 압축하였다. 내부 기체의 압력은 어떻게 되는가?',
     [('3배가 된다','같은 수의 분자가 더 좁은 공간에서 벽에 충돌하는 횟수가 늘기 때문',True,None),
      ('1/3로 작아진다','부피와 압력이 같은 비율로 함께 변하기 때문',False,'정비례 오개념'),
      ('변하지 않는다','압력은 기체의 고유한 성질이라 일정하기 때문',False,'압력 불변 오개념'),
      ('2/3로 작아진다','줄어든 부피만큼 압력이 빠져나가기 때문',False,'수학적 합산 오류')]),
   F('객관식 2 — 그래프 고르기',
     '일정한 온도에서 부피(가로축)에 따른 압력(세로축)의 관계를 가장 잘 나타낸 그래프는?',
     [('g_inv','A',True),('g_prop','B',False),('g_flat','C',False),('g_lin','D',False)],
     [('A: 압력과 부피의 곱이 일정하여, 부피가 늘면 압력이 반비례로 줄기 때문',True,None),
      ('B: 부피와 압력이 같은 비율로 함께 커지기 때문',False,'정비례 오개념'),
      ('C: 밀폐되어 있으면 부피가 변해도 압력이 일정하기 때문',False,'압력 불변 오개념'),
      ('D: 부피가 일정량 늘 때마다 압력이 일정량씩 줄기 때문',False,'선형 감소 오개념')]),
 ]),
 ('오개념 B2 — 압축·팽창 시 입자 보존', [
   F('객관식 1 — 입자 그림 고르기(크기·개수·빈 공간)',
     '아래는 분자 8개가 든 주사기의 처음 상태(부피 V)이다. 이 기체를 절반 부피로 압축하였을 때의 모습으로 가장 옳은 그림은?',
     [('p_same','A',True),('p_small','B',False),('p_four','C',False),('p_filled','D',False)],
     [('A: 분자의 수와 크기는 그대로이고, 분자 사이의 빈 공간(간격)만 좁아지기 때문',True,None),
      ('B: 압력이 분자를 눌러 작게 만들기 때문',False,'입자 크기 축소 오개념'),
      ('C: 좁아진 공간에 다 들어가도록 분자 일부가 사라지기 때문',False,'입자 수 감소 오개념'),
      ('D: 압축하면 입자 사이가 무언가로 채워져 빈틈이 없어지기 때문',False,'연속체·빈 공간 부정 오개념')],
     before='p_before'),
   T('객관식 2 — 실제 상황(팽창 시 질량)',
     '헬륨을 채운 풍선이 하늘로 올라가면서 점점 부풀어 커진다(헬륨은 새지 않는다). 이때 풍선 속 기체의 질량은 어떻게 되는가?',
     [('변하지 않는다','헬륨이 새지 않으므로 분자 수와 질량은 보존되고, 변하는 것은 부피(밀도)뿐이기 때문',True,None),
      ('줄어든다','위로 올라가 가벼워졌으니 질량도 줄기 때문',False,'부력·질량 혼동'),
      ('늘어난다','부풀어 커졌으니 기체의 양도 늘기 때문',False,'부피·질량 혼동'),
      ('부피에 비례해 변한다','부피가 변하면 질량도 따라 변하기 때문',False,'부피·질량 혼동')]),
 ]),
 ('오개념 B3 — 진공 흡입', [
   T('객관식 1 — 직접 시나리오(주사기 당김)',
     '끝이 막힌 주사기의 피스톤을 바깥으로 당겨 내부 부피가 팽창하였다. 이때 내부 압력은 어떻게 되는가?',
     [('감소한다','공간이 넓어져 단위 부피당 벽에 부딪치는 횟수가 줄기 때문',True,None),
      ('증가한다','당길 때 생긴 진공이 입자를 바깥으로 끌어당기기 때문',False,'진공 흡입 오개념'),
      ('변하지 않는다','밀폐되어 있어 압력이 일정하기 때문',False,'압력 불변 오개념'),
      ('0이 된다','빈 공간이 되어 압력이 없어지기 때문',False,'진공 오개념')]),
   T('객관식 2 — 실제 상황(빨대)',
     '빨대로 음료를 마실 때 음료가 위로 올라오는 이유로 가장 옳은 것은?',
     [('입 안의 압력이 낮아져, 바깥 대기압이 음료를 밀어 올린다','입 안 압력이 대기압보다 낮아진 압력 차가 음료를 밀어 올리기 때문',True,None),
      ('입 안에 생긴 진공이 음료를 빨아올린다','진공이 액체를 끌어당기는 힘을 내기 때문',False,'진공 흡입 오개념'),
      ('빨대 속 빈 공간이 음료를 끌어당긴다','빈 공간 자체가 힘의 원천이기 때문',False,'빈 공간 인력 오개념'),
      ('빨대가 음료를 끌어올린다','빨대 재질이 액체를 당기기 때문',False,'비압력 오개념')]),
 ]),
 ('오개념 D1 — 돌턴 법칙 기본(분포와 독립성)', [
   F('객관식 1 — 혼합 기체 분포 그림 고르기',
     '헬륨(○)과 이산화탄소(●)를 같은 양으로 한 용기에 넣고 충분히 둔 뒤의 분포로 가장 옳은 그림은? (○ 가벼운 헬륨, ● 무거운 이산화탄소, 각 6개)',
     [('m_homog','A',True),('m_layer','B',False),('m_split','C',False),('m_cluster','D',False)],
     [('A: 분자가 끊임없이 무작위로 운동·확산하여 무게나 종류와 상관없이 고르게 섞이기 때문',True,None),
      ('B: 무거운 기체는 가라앉고 가벼운 기체는 떠올라 위·아래로 갈리기 때문',False,'완전 층화 오개념'),
      ('C: 두 기체가 용기를 영역으로 나누어 따로 차지하기 때문',False,'부피 분할 오개념'),
      ('D: 같은 종류의 분자끼리 서로 끌어당겨 뭉치기 때문',False,'동종 인력 오개념')]),
   T('객관식 2 — 분압의 독립성',
     '산소만 들어 있어 산소의 부분 압력이 1기압인 통에, 질소 기체를 추가로 더 넣었다. 산소의 부분 압력은 어떻게 되는가?',
     [('변하지 않고 1기압 그대로이다','각 기체는 다른 기체가 있든 없든 자기 부분 압력을 그대로 유지하기 때문',True,None),
      ('줄어든다','질소가 들어와 산소가 차지할 몫이 줄기 때문',False,'분압 분산 오개념'),
      ('늘어난다','전체 압력이 올라가니 산소 압력도 따라 오르기 때문',False,'전체압 혼동'),
      ('절반으로 줄어든다','두 기체가 압력을 똑같이 나누어 가지기 때문',False,'압력 분할 오개념')]),
 ]),
 ('오개념 D2 — 압력과 분자량의 관계', [
   T('객관식 1 — 직접 판단(분자량과 부분 압력)',
     '수소 1몰과 이산화탄소 1몰을 같은 용기에 넣어 섞었다(이산화탄소가 약 22배 무겁다). 두 기체의 부분 압력을 비교하면?',
     [('두 기체의 부분 압력이 같다','압력은 분자의 무게가 아니라 분자의 수(몰수)로 정해지기 때문',True,None),
      ('이산화탄소가 훨씬 크다','무거운 분자가 벽에 더 강하게 충돌하기 때문',False,'분자량 오개념'),
      ('수소가 더 크다','가벼워서 더 빠르고 자주 충돌하기 때문',False,'분자량·속도 혼동'),
      ('비교할 수 없다','두 기체의 압력은 종류가 다르기 때문',False,'비교 불가 오개념')]),
   T('객관식 2 — 충돌·시간 경과 후 부분 압력',
     '수소와 이산화탄소를 한 용기에 섞었다. 분자들이 서로 충돌하며 충분한 시간이 지난 뒤, 각 기체의 부분 압력은 어떻게 되는가?',
     [('충돌 전과 똑같이 각자 그대로 유지된다','충돌은 분자끼리 에너지를 주고받을 뿐, 각 기체의 부분 압력은 제 몰수와 온도로 정해져 변하지 않기 때문',True,None),
      ('무거운 이산화탄소의 부분 압력이 더 커진다','무거운 분자가 충돌로 압력을 더 많이 차지하기 때문',False,'분자량 오개념'),
      ('두 기체의 부분 압력이 같은 값으로 같아진다','충돌하면서 압력이 고르게 나누어지기 때문',False,'압력 평준화 오개념'),
      ('가벼운 수소의 부분 압력이 더 커진다','가볍고 빨라 벽에 더 자주 부딪치기 때문',False,'속도·압력 혼동')]),
 ]),
 ('오개념 G1 — 압력의 본질', [
   T('객관식 1 — 현상 설명',
     '일정한 온도의 밀폐된 강철 통 속 기체가 벽을 누르는 가장 근본적인 이유는?',
     [('분자가 끊임없이 무작위로 움직이며 벽에 충돌하기 때문','충돌할 때마다 운동량이 전달되고, 그 힘의 총합이 압력이기 때문',True,None),
      ('기체가 부풀어 오르려는 힘 때문','기체에는 넓어지려는 성질이 들어 있기 때문',False,'내재적 팽창력 오개념'),
      ('기체의 무게가 벽을 누르기 때문','무거울수록 강하게 눌러 압력이 곧 무게이기 때문',False,'압력=무게 오개념'),
      ('분자끼리 밀어내 벽까지 밀리기 때문','분자 사이의 미는 힘이 벽으로 전달되기 때문',False,'입자 반발 오개념')]),
   T('객관식 2 — 실제 상황(무중력)',
     '지구의 실험실과 무중력 상태의 우주 정거장에 같은 용기·온도·양의 기체가 있다. 두 곳에서 벽이 받는 압력을 비교하면?',
     [('두 곳이 사실상 같다','압력은 분자의 무작위 운동에서 나오고 그 운동은 온도가 정하므로 중력과 무관하기 때문',True,None),
      ('무중력에서는 거의 0이 된다','압력의 근원이 무게라서 중력이 없으면 사라지기 때문',False,'압력=무게 오개념'),
      ('무중력에서 더 높다','중력이 없으면 분자가 더 자유로워 세게 밀기 때문',False,'운동 오인 오개념'),
      ('지구에서 더 높다','기체의 무게가 압력에 더해지기 때문',False,'압력=무게 오개념')]),
 ]),
 ('오개념 G2 — 압력과 무관한 변수(시간·중력·재질)', [
   T('객관식 1 — 직접 판단(변수: 시간)',
     '온도가 일정한 방에 놓인 밀폐 용기를 건드리지 않고 사흘이 지났다. 내부 압력은 어떻게 되는가?',
     [('변하지 않는다','온도·부피·기체의 양이 그대로이면 분자의 평균 운동이 유지되어 압력이 일정하기 때문',True,None),
      ('분자가 가라앉아 낮아진다','가만히 두면 무거운 분자가 바닥에 가라앉기 때문',False,'중력 침강 오개념'),
      ('분자가 지쳐 느려져 낮아진다','시간이 흐르면 분자가 운동 에너지를 잃기 때문',False,'입자 운동 소멸 오개념'),
      ('충돌이 쌓여 높아진다','충돌이 거듭될수록 미는 힘이 커지기 때문',False,'충돌 누적 오개념')]),
   T('객관식 2 — 실제 상황(변수: 중력·방향, 등방성)',
     '보통 크기의 밀폐 용기 속 기체가 있다. 용기의 윗면과 아랫면이 받는 압력을 비교하면?',
     [('사실상 같다','분자가 모든 방향으로 무작위 운동하여 모든 면에 같은 크기로 작용하기 때문',True,None),
      ('아랫면이 뚜렷이 크다','기체의 무게가 아래로 눌리기 때문',False,'중력 쏠림 오개념'),
      ('윗면이 크다','가벼운 분자가 위로 떠오르기 때문',False,'부력 오개념'),
      ('위·아래에는 압력이 없다','압력은 옆으로 미는 힘이기 때문',False,'작용면 오개념')]),
 ]),
]

# ---------------- build ----------------
doc = Document()
for sname in ('Normal',):
    st = doc.styles[sname]; st.font.name = KR; st.font.size = Pt(10)
    rpr = st.element.get_or_add_rPr(); rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rpr.append(rf)
    rf.set(qn('w:eastAsia'), KR); rf.set(qn('w:ascii'), KR); rf.set(qn('w:hAnsi'), KR)

# page setup: 여백·머리말·꼬리말(쪽번호)·표지 분리
sec = doc.sections[0]
sec.top_margin = Inches(0.8); sec.bottom_margin = Inches(0.8)
sec.left_margin = Inches(0.7); sec.right_margin = Inches(0.7)
sec.header_distance = Inches(0.4); sec.footer_distance = Inches(0.4)
sec.different_first_page_header_footer = True
_hp = sec.header.paragraphs[0]; _hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_kr(_hp.add_run('기체 개념 진단 검사지 · 작성자용'), size=8.5, color=GREY)
_fp = sec.footer.paragraphs[0]; _fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_kr(_fp.add_run('— '), size=9, color=GREY); add_page_number(_fp); set_kr(_fp.add_run(' —'), size=9, color=GREY)

# title
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('기체 개념 진단 검사지'); set_kr(r, size=20, bold=True, color=NAVY)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('핵심 오개념 7개 (보일 3 · 돌턴 2 · 기체 일반 2) · 진위형 7 + 객관식 14'); set_kr(r, size=10.5, color=RGBColor(0x6b,0x72,0x80))
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('※ 작성자용 — 정답(✓)·진단 오개념 표시 포함. 실제 배포본에서는 숨김.'); set_kr(r, size=9, color=RED)
doc.add_paragraph()
_ct = doc.add_table(rows=0, cols=2); _ct.alignment = WD_TABLE_ALIGNMENT.CENTER
style_table(_ct, [1.7, 3.3])
for _k, _v in [('검사 영역', '보일 · 돌턴 · 기체 일반'), ('핵심 오개념', '7개'),
               ('문항 수', '진위형 7 + 객관식 14'), ('형식', '삼층 진단(답 · 이유 · 확신 정도)'),
               ('작성일', '2026-05-21')]:
    _cs = _ct.add_row().cells
    cell_text(_cs[0], _k, size=10, bold=True); shade(_cs[0], 'F3F4F6')
    cell_text(_cs[1], _v, size=10)
doc.add_page_break()

def heading(text, size=14, color=NAVY):
    p = doc.add_paragraph(); r = p.add_run(text); set_kr(r, size=size, bold=True, color=color)
    pf = p.paragraph_format; pf.space_before = Pt(2); pf.space_after = Pt(5); pf.keep_with_next = True
    return p

# ---- 문항 구성 개요 (진위형 ↔ 객관식 연결) ----
heading('문항 구성 — 진위형과 객관식의 연결', size=15)
_pp = doc.add_paragraph()
set_kr(_pp.add_run('핵심 오개념 7개. 각 오개념마다 진위형 1문항(빠른 선별)과 객관식 2문항(답·이유·확신 정도)이 짝을 이룬다. '
                   '진위형은 객관식과 서로 다른 측면을 물어 정답 단서 노출을 피한다. 총 진위형 7 + 객관식 14 = 21문항.'), size=10, color=MUTED)
ov = doc.add_table(rows=1, cols=5); ov.style = 'Table Grid'
style_table(ov, [0.5, 1.55, 2.35, 2.4, 0.45]); repeat_header(ov.rows[0])
for _c, _t in zip(ov.rows[0].cells, ['코드', '오개념', '진위형 (정답)', '객관식 ① · ② (측면)', '계']):
    cell_text(_c, _t, size=9.3, bold=True); shade(_c, 'F3F4F6')
for (sect_title, qs), (area, stmt, ans) in zip(SECTIONS, TF):
    code = sect_title.split(' ')[1]
    name = sect_title.split('— ')[1]
    a1 = qs[0]['ang'].split('— ')[-1]
    a2 = qs[1]['ang'].split('— ')[-1]
    cells = ov.add_row().cells
    cell_text(cells[0], code, size=9.5, bold=True, color=NAVY)
    cell_text(cells[1], name, size=8.8)
    _p2 = cells[2].paragraphs[0]
    set_kr(_p2.add_run(stmt + ' '), size=8.6)
    set_kr(_p2.add_run('(' + ans + ')'), size=8.6, bold=True, color=(GREEN if ans == '참' else RED))
    cells[3].text = ''
    _pa = cells[3].paragraphs[0]; set_kr(_pa.add_run('① ' + a1), size=8.6)
    _pb = cells[3].add_paragraph(); set_kr(_pb.add_run('② ' + a2), size=8.6)
    cell_text(cells[4], '3', size=9.5)
doc.add_page_break()

# 진위형
heading('진위형 7문항', size=15)
p = doc.add_paragraph(); r = p.add_run('진술이 과학적으로 맞으면 「참」, 틀리면 「거짓」. 참 3 / 거짓 4. 객관식과는 서로 다른 측면을 묻는다.'); set_kr(r, size=9, color=MUTED)
t = doc.add_table(rows=1, cols=3); t.style = 'Table Grid'
style_table(t, [1.1, 4.6, 0.9]); repeat_header(t.rows[0])
hdr = t.rows[0].cells
for c, txt, wd in zip(hdr, ['영역','진술','정답'], [9,9,9]):
    cell_text(c, txt, size=9.5, bold=True); shade(c, 'F3F4F6')
for area, stmt, ans in TF:
    cells = t.add_row().cells
    cell_text(cells[0], area, size=9)
    cell_text(cells[1], stmt, size=9.5)
    cell_text(cells[2], ans, size=9.5, bold=True, color=(GREEN if ans=='참' else RED))
doc.add_paragraph()

def add_reason_table(reasons):
    t = doc.add_table(rows=0, cols=1); t.style = 'Table Grid'
    style_table(t, [6.4])
    for txt, correct, code in reasons:
        cells = t.add_row().cells
        p = cells[0].paragraphs[0]
        run = p.add_run(('✓ ' if correct else '') + txt)
        set_kr(run, size=9.5, bold=correct, color=(GREEN if correct else None))
        if code:
            cr = p.add_run('   [' + code + ']'); set_kr(cr, size=8.5, color=RED)
        if correct: shade(cells[0], GREENFILL_STRONG)
    doc.add_paragraph()

def add_answer_reason_table(rows):
    t = doc.add_table(rows=1, cols=2); t.style = 'Table Grid'
    style_table(t, [1.7, 4.7]); repeat_header(t.rows[0])
    h = t.rows[0].cells
    cell_text(h[0], '답', size=9.5, bold=True); shade(h[0], 'F3F4F6')
    cell_text(h[1], '이유', size=9.5, bold=True); shade(h[1], 'F3F4F6')
    for ans, reason, correct, code in rows:
        cells = t.add_row().cells
        p0 = cells[0].paragraphs[0]
        ra = p0.add_run(('✓ ' if correct else '') + ans); set_kr(ra, size=9.5, bold=correct, color=(GREEN if correct else None))
        p1 = cells[1].paragraphs[0]
        rr = p1.add_run(reason); set_kr(rr, size=9.5, color=(GREEN if correct else None))
        if code:
            cr = p1.add_run('  [' + code + ']'); set_kr(cr, size=8.5, color=RED)
        if correct:
            shade(cells[0], GREENFILL_STRONG); shade(cells[1], GREENFILL_SOFT)
    doc.add_paragraph()

def add_fig_choices(figs):
    t = doc.add_table(rows=2, cols=len(figs)); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    fixed_table(t); set_col_widths(t, [6.4/len(figs)]*len(figs))
    for i, (fid, label, correct) in enumerate(figs):
        cell = t.cell(0, i)
        pr = cell.paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pr.add_run().add_picture(make_fig(fid), height=Inches(0.95))
        lc = t.cell(1, i); lp = lc.paragraphs[0]; lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = lp.add_run(label + (' ✓' if correct else '')); set_kr(lr, size=10, bold=True, color=(GREEN if correct else None))
    doc.add_paragraph()

for _idx, (sect_title, qs) in enumerate(SECTIONS):
    if _idx > 0:
        doc.add_page_break()
    heading(sect_title, size=14)
    for q in qs:
        p = doc.add_paragraph(); r = p.add_run(' ' + q['ang'] + ' '); set_kr(r, size=10, bold=True, color=RGBColor(0xFF,0xFF,0xFF))
        para_shade(p, 'D97706'); p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
        p = doc.add_paragraph(); r = p.add_run(q['stem']); set_kr(r, size=11)
        p.paragraph_format.keep_with_next = True; p.paragraph_format.space_after = Pt(4)
        if q['kind'] == 'text':
            add_answer_reason_table(q['rows'])
        else:
            if q.get('before'):
                bp = doc.add_paragraph()
                bp.add_run().add_picture(make_fig(q['before']), height=Inches(1.2))
                cap = bp.add_run('   ← 처음 상태 (부피 V, 분자 8개)'); set_kr(cap, size=9, color=RGBColor(0x6b,0x72,0x80))
            add_fig_choices(q['figs'])
            add_reason_table(q['reasons'])

doc.add_paragraph()
fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run('기체 개념 진단 검사지 · 작성자용 정본 · 2026-05-21'); set_kr(fr, size=8.5, color=RGBColor(0x9c,0xa3,0xaf))

out = r'C:\Users\user\project\pchem-lab-project\measure\검사지-최종문항.docx'
doc.save(out)
print('saved:', out)
